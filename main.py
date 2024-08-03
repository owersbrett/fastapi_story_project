from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import List
from fastapi.responses import FileResponse
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

app = FastAPI()

class Choice(BaseModel):
    text: str
    next_page: int

class Page(BaseModel):
    number: int
    content: str
    choices: List[Choice]

class Story(BaseModel):
    title: str
    pages: List[Page]

# In-memory storage for our story
story_database = {}

@app.post("/import_story")
async def import_story(story: Story):
    story_database[story.title] = story
    return {"message": f"Story '{story.title}' imported successfully"}

@app.get("/export/story/{title}")
async def export_story(title: str):
    if title not in story_database:
        raise HTTPException(status_code=404, detail="Story not found")
    
    story = story_database[title]
    doc = Document()
    doc.add_heading(story.title, 0)
    
    for page in story.pages:
        # Create a bookmark for each page
        # p = doc.add_paragraph()
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(page.number))
        bookmark_start.set(qn("w:name"), f"Page{page.number}")
        # p._element.append(bookmark_start)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(page.number))
        # p._element.append(bookmark_end)
        
        heading = doc.add_heading(f'Page {page.number}', level=1)
        heading._element.append(bookmark_start)
        heading._element.append(bookmark_end)
        doc.add_paragraph(page.content)
       
        if page.choices:
            
            choices = doc.add_paragraph("")
            insertHR(choices)
            for choice in page.choices:
                p_choice = doc.add_paragraph(style='ListBullet')
                run = p_choice.add_run(f"{choice.text} (Go to page {choice.next_page})")
                # Create hyperlink to the corresponding page
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("w:anchor"), f"Page{choice.next_page}")
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rStyle = OxmlElement("w:rStyle")
                rStyle.set(qn("w:val"), "Hyperlink")
                rPr.append(rStyle)
                r.append(rPr)
                r_text = OxmlElement("w:t")
                r_text.text = run.text
                r.append(r_text)
                hyperlink.append(r)
                p_choice._element.clear_content()
                p_choice._element.append(hyperlink)
                
        
        doc.add_page_break()
    
    filename = f"{title.replace(' ', '_')}.docx"
    filepath = f"/tmp/{filename}"
    doc.save(filepath)
    
    return FileResponse(filepath, filename=filename)

# Example usage
example_story = Story(
    title="Cave Adventure",
    pages=[
        Page(
            number=1,
            content="You wake up in a dimly lit cave, the air damp and musty.",
            choices=[
            ]
        ),
        Page(
            number=2,
            content="As your eyes adjust, you notice two tunnels leading out of the chamber. A flickering torch on the wall casts dancing shadows, revealing ancient symbols etched into the stone floor.",
            choices=[
                Choice(text="Take the left tunnel", next_page=3),
                Choice(text="Take the right tunnel", next_page=4)
            ]
        ),
        Page(
            number=3,
            content="The left tunnel narrows as you proceed, the ceiling dropping lower with each step. Suddenly, you hear a low rumble. Before you can react, the tunnel collapses, trapping you in total darkness. Your adventure ends here.",
            choices=[]
        ),
        Page(
            number=4,
            content="The right tunnel opens into a vast underground cavern. A subterranean river cuts through the center, its rushing waters echoing off the walls. Three rope bridges span the chasm, each leading to a different opening on the far side.",
            choices=[
                Choice(text="Cross the left bridge", next_page=5),
                Choice(text="Cross the middle bridge", next_page=6),
                Choice(text="Cross the right bridge", next_page=7)
            ]
        ),
        Page(
            number=5,
            content="As you carefully make your way across the left bridge, you notice strange, glowing fungi growing on the cavern walls. Their soft blue light illuminates a hidden path behind a boulder. You decide to investigate and discover an ancient treasure room filled with golden artifacts.",
            choices=[
                Choice(text="Take some treasure and continue", next_page=8),
                Choice(text="Leave the treasure and explore further", next_page=9)
            ]
        ),
        Page(
            number=6,
            content="The middle bridge creaks ominously as you step onto it. Halfway across, you encounter a wise old hermit who offers you a choice: a map of the caves or a magical amulet.",
            choices=[
                Choice(text="Choose the map", next_page=10),
                Choice(text="Choose the amulet", next_page=11)
            ]
        ),
        Page(
            number=7,
            content="The right bridge seems the sturdiest of the three. As you cross, you notice a series of intricate carvings along the cavern wall, depicting an ancient civilization. At the end of the bridge, you find a lever embedded in the rock face.",
            choices=[
                Choice(text="Pull the lever", next_page=12),
                Choice(text="Ignore the lever and continue", next_page=13)
            ]
        )
    ]
)
@app.on_event("startup")
async def startup_event():
    await import_story(example_story)

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)